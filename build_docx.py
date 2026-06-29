"""从 Markdown 简历生成可编辑 Word 文档 — 与 HTML 样式一致"""
import re, sys, io
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
md_text = Path('data/resume_optimized.md').read_text('utf-8')

# Colors
DARK_BLUE = RGBColor(0x1a, 0x1a, 0x2e)
MID_BLUE  = RGBColor(0x16, 0x21, 0x3e)
DEEP_BLUE = RGBColor(0x0f, 0x34, 0x60)
ACCENT    = RGBColor(0xe9, 0x45, 0x60)
GREEN     = RGBColor(0x5a, 0x9a, 0x5a)
BLACK     = RGBColor(0x1a, 0x1a, 0x1a)
GRAY      = RGBColor(0x55, 0x55, 0x55)
LGRAY     = RGBColor(0x88, 0x88, 0x88)
WHITE     = RGBColor(0xff, 0xff, 0xff)
LIGHT_BG  = RGBColor(0xf8, 0xf9, 0xfa)

def get_section(title):
    pattern = rf'## {title}\n(.*?)(?=\n## |\n---|\Z)'
    m = re.search(pattern, md_text, re.DOTALL)
    return m.group(1).strip() if m else ''

def add_blue_header(doc):
    """Add the dark blue header section with name, title, photo placeholder"""
    # Create a table for the header
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table width to full page
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # Left cell - name and title
    left = table.cell(0, 0)
    left.width = Cm(12)
    # Blue background
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1a1a2e')
    shading.set(qn('w:val'), 'clear')
    left._tc.get_or_add_tcPr().append(shading)

    left.paragraphs[0].clear()
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(1.2)

    target_match = re.search(r'目标方向[：:]\s*(.+)', md_text)
    target = target_match.group(1).strip() if target_match else 'AIGC内容创作者'

    run = p.add_run(target)
    run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = WHITE

    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.left_indent = Cm(1.2)
    run2 = p2.add_run('AIGC内容创作 · AI科技自媒体 · 深圳 · 15-20K')
    run2.font.size = Pt(10); run2.font.color.rgb = RGBColor(0xcc, 0xcc, 0xdd)

    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(16)

    # Right cell - photo placeholder
    right = table.cell(0, 1)
    right.width = Cm(6)
    shading2 = OxmlElement('w:shd')
    shading2.set(qn('w:fill'), '16213e')
    shading2.set(qn('w:val'), 'clear')
    right._tc.get_or_add_tcPr().append(shading2)

    right.paragraphs[0].clear()
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr.paragraph_format.space_before = Pt(14)
    run_pr = pr.add_run('证件照')
    run_pr.font.size = Pt(10); run_pr.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    doc.add_paragraph()  # spacer

def add_section_title(doc, title):
    """Add a section title with green left border"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    # Add green left border via paragraph shading/border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '12')
    left_border.set(qn('w:space'), '6')
    left_border.set(qn('w:color'), '5a9a5a')
    pBdr.append(left_border)
    pPr.append(pBdr)

    run = p.add_run(title)
    run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = BLACK
    return p

def add_paragraph(doc, text, size=10, bold=False, color=BLACK, indent=0, spacing_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(spacing_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold; run.font.color.rgb = color
    return p

def add_bullet(doc, text, size=10, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run('• ' + text)
    run.font.size = Pt(size); run.font.color.rgb = color
    return p

# ━━━━ Build Document ━━━━
doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(1.0)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# Set default font
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# ── Header ──
add_blue_header(doc)

# ── Summary ──
add_section_title(doc, '个人概述')
summary = get_section('个人简介')
add_paragraph(doc, summary, size=10, color=GRAY)

# ── Skills ──
add_section_title(doc, '核心技能')
for line in get_section('技能栈').split('\n'):
    m = re.match(r'- \*\*(.+?)\*\*: (.+)', line)
    if m:
        level = m.group(1)
        items = m.group(2)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run_l = p.add_run(f'{level}: ')
        run_l.font.size = Pt(10); run_l.font.bold = True; run_l.font.color.rgb = BLACK
        run_i = p.add_run(items)
        run_i.font.size = Pt(10); run_i.font.color.rgb = DARK_BLUE

# ── Work Experience ──
add_section_title(doc, '工作经历')
exp_text = get_section('工作经历')
for block in re.split(r'\n### ', '\n' + exp_text):
    if not block.strip(): continue
    lines = block.strip().split('\n')
    h = lines[0].replace('### ','')
    dur = ''
    for l in lines[1:]:
        s = l.strip()
        if s.startswith('*'):
            dur = s.replace('*','').strip()
            break

    add_paragraph(doc, h, size=12, bold=True, spacing_after=2)
    if dur:
        add_paragraph(doc, dur, size=9, color=GREEN, spacing_after=4)

    for l in lines[1:]:
        s = l.strip()
        if s.startswith('- ') and '技术栈' not in s:
            add_bullet(doc, s[2:], size=10)
    add_paragraph(doc, '', size=4, spacing_after=2)  # spacer

# ── Projects ──
add_section_title(doc, '项目经验')
proj_text = get_section('项目经验')
for block in re.split(r'\n### ', '\n' + proj_text):
    if not block.strip(): continue
    lines = block.strip().split('\n')
    name = lines[0].replace('### ','')
    role = ''
    for l in lines[1:]:
        s = l.strip()
        if s.startswith('*'):
            role = s.replace('*','').strip()
            break

    header = name
    if role: header += f'  —  {role}'
    add_paragraph(doc, header, size=12, bold=True, spacing_after=4)

    for l in lines[1:]:
        s = l.strip()
        if s.startswith('- ') and '技术栈' not in s:
            add_bullet(doc, s[2:], size=10)
    add_paragraph(doc, '', size=4, spacing_after=2)

# ── Education ──
add_section_title(doc, '教育背景')
edu_text = get_section('教育背景')
edu_lines = edu_text.split('\n')
parts = edu_lines[0].replace('**','').split('|')
edu_school = parts[0].strip() if len(parts)>0 else ''
edu_degree = parts[1].strip() if len(parts)>1 else ''
edu_major = parts[2].strip() if len(parts)>2 else ''
add_paragraph(doc, f'{edu_school}  |  {edu_degree}  |  {edu_major}', size=11, bold=True)
for l in edu_lines[1:]:
    if l.strip().startswith('-'):
        add_bullet(doc, l.strip('- ').strip(), size=10, color=GRAY)

# ── Portfolio ──
add_section_title(doc, '作品集 & 社区影响力')
portfolio_text = get_section('作品集')
for l in portfolio_text.split('\n'):
    if l.strip().startswith('-'):
        # Remove URLs for clean text
        text = re.sub(r'https?://[^\s)]+', '', l.strip('- ')).strip()
        if text:
            add_paragraph(doc, text, size=10, color=GRAY, spacing_after=2)

# ── Job Target ──
add_section_title(doc, '求职意向')
job_text = get_section('求职意向')
for l in job_text.split('\n'):
    s = l.strip('- ').strip()
    if s:
        add_paragraph(doc, s, size=10, color=BLACK, spacing_after=2)

# Footer
add_paragraph(doc, '', size=4, spacing_after=8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('叁十三 · AIGC创作者  |  深圳  |  sanshisanAIGC')
run.font.size = Pt(8); run.font.color.rgb = LGRAY

# ━━━━ Save ━━━━
output = Path('data/resume.docx')
doc.save(str(output))
print(f'Word 简历已生成: {output}')
print(f'A4 尺寸 | 样式与HTML一致 | 可直接编辑')
